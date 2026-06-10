#!/usr/bin/env python3
"""
md_to_docx.py — 教材 Markdown → Word 转换工具（v3，python-docx + OMML）

用 python-docx 构建 docx，用 latex_to_omml 将 LaTeX 公式转为 Word 可编辑的 OMML 方程。

支持：
- LaTeX $$ 公式 → 可双击编辑的 Word 公式
- 行内 $ 公式 → 行内 Word 公式
- Markdown 表格 → 原生 Word 表格
- Mermaid 代码块 → 描述文字 + 源码区块
- 标题层级 (#/##/### → Heading1/2/3)
- 图片 ![alt](path) → 嵌入图片（如文件存在）

依赖: python-docx, lxml, latex_to_omml.py

用法:
    python3 scripts/md_to_docx.py single 第N章.md -o 第N章.docx
    python3 scripts/md_to_docx.py dir output/ -o 全书.docx
"""

import argparse
import os
import re
import sys
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import docx.oxml

# ── 导入 OMML 转换器 ──
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__))))
from latex_to_omml import latex_to_omml
from lxml import etree

# ═══════════════════════════════════════════════════
# OMML namespace
# ═══════════════════════════════════════════════════

M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# ═══════════════════════════════════════════════════
# 公式预处理
# ═══════════════════════════════════════════════════


def clean_latex(formula: str) -> str:
    """清理 LaTeX 公式，使其可被 latex_to_omml.py 解析。"""
    # 1) 移除 \tag{N-M} 行
    formula = re.sub(r'^\\tag\{[\d-]+\}\s*$', '', formula, flags=re.MULTILINE)
    formula = re.sub(r'^\\\\tag\{[\d-]+\}\s*$', '', formula, flags=re.MULTILINE)
    formula = re.sub(r'\n\\tag\{[\d-]+\}\s*\n', '\n', formula)
    # 2) \xrightarrow{a} → a \to
    formula = re.sub(r'\\xrightarrow\{([^}]*)\}', r'\1 \\to', formula)
    # 3) \displaystyle / \limits 移除
    formula = re.sub(r'\\displaystyle\s*', '', formula)
    formula = re.sub(r'\\limits\s*', '', formula)
    # 4) \begin{aligned} / \end{aligned} 保留（latex_to_omml 支持）
    return formula.strip()


# ═══════════════════════════════════════════════════
# Markdown 解析
# ═══════════════════════════════════════════════════

def parse_blocks(md: str) -> list[dict]:
    """将 markdown 解析为块列表：{type, content, ...}"""
    blocks = []

    # 剥离 YAML frontmatter
    md = re.sub(r'^---\n.*?\n---\n', '', md, flags=re.DOTALL)

    # 按段落拆分，保留代码块和表格
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Mermaid 代码块 ──
        if line.strip().startswith('```mermaid'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append({
                'type': 'mermaid',
                'code': '\n'.join(code_lines),
            })
            continue

        # ── 其他代码块 ──
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append({
                'type': 'code',
                'code': '\n'.join(code_lines),
            })
            continue

        # ── 表格（连续 | 行） ──
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_rows.append(lines[i].strip())
                i += 1
            if len(table_rows) >= 2:
                blocks.append({
                    'type': 'table',
                    'rows': table_rows,
                })
            continue

        # ── 空行 ──
        if not line.strip():
            i += 1
            continue

        # ── 普通段落（含公式） ──
        para_lines = []
        while i < len(lines) and lines[i].strip():
            if lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                break
            if lines[i].strip().startswith('```'):
                break
            para_lines.append(lines[i])
            i += 1

        if para_lines:
            text = '\n'.join(para_lines)
            blocks.append({
                'type': 'paragraph',
                'text': text,
            })

    return blocks


# ═══════════════════════════════════════════════════
# 构建 docx
# ═══════════════════════════════════════════════════


def heading_level(text: str):
    """检测文本是否以 # 开头，返回 (级别, 内容) 或 (None, text)。"""
    m = re.match(r'^(#{1,3})\s+(.*)', text)
    if m:
        return len(m.group(1)), m.group(2).strip()
    return None, text


def add_formula_paragraph(doc: Document, latex: str):
    """将 LaTeX 公式作为 OMML 段落插入 docx。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 创建 oMathPara 元素
    oMathPara = etree.SubElement(
        p._element, qn('m:oMathPara')
    )
    # 创建 oMath 元素
    try:
        latex_clean = clean_latex(latex)
        oMath = latex_to_omml(latex_clean)
        oMathPara.append(oMath)
    except Exception as e:
        # fallback: 显示原始 LaTeX
        run = p.add_run(f"[公式解析失败: {latex[:60]}]")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    return p


def add_inline_formula(paragraph, latex: str):
    """在段落中插入行内 OMML 公式。"""
    try:
        latex_clean = clean_latex(latex)
        oMath = latex_to_omml(latex_clean)
        # 插入到段落最后
        paragraph._element.append(oMath)
    except Exception:
        paragraph.add_run(f"[{latex[:30]}]")


def get_horizontal_alignment(table_align: str) -> str:
    """Helper to return horizontal alignment string for table cells."""
    return table_align


def add_table(doc: Document, rows_text: list[str]):
    """将 markdown 表格转为 Word 表格。"""
    # 解析 rows
    parsed = []
    for row_text in rows_text:
        cells = [c.strip() for c in row_text.strip('|').split('|')]
        parsed.append(cells)

    if len(parsed) < 2:
        return

    header = parsed[0]
    # 第二行是分隔符（|---|），确定对齐方式
    alignments = ['left'] * len(header)
    if len(parsed) > 1:
        sep = parsed[1]
        for j, s in enumerate(sep):
            s_stripped = s.strip()
            if s_stripped.startswith(':') and s_stripped.endswith(':'):
                alignments[j] = 'center'
            elif s_stripped.endswith(':'):
                alignments[j] = 'right'
            elif s_stripped.startswith(':'):
                alignments[j] = 'left'

    data_rows = parsed[2:] if len(parsed) > 2 else []

    if not data_rows:
        return

    n_cols = len(header)
    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.style = 'Table Grid'

    # 表头
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for i, row in enumerate(data_rows):
        for j, cell_text in enumerate(row):
            if j >= n_cols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            # 处理行内公式
            process_inline_math(p, cell_text)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(10)

    return table


def process_markdown_text(paragraph, text: str):
    """处理段落中的 Markdown 语法：**粗体**、*斜体*、行内 $公式$、`代码`"""
    # 用正则分段，保留分隔符以便区分不同类型的标记
    # 优先级: $$ > $ > ** > * > ` > [text](url) > plain
    pattern = r'(\$\$[\s\S]*?\$\$|\$[^$]+\$|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[([^\]]+)\]\(([^)]+)\))'

    parts = re.split(pattern, text)
    i = 0
    while i < len(parts):
        part = parts[i]
        if not part:
            i += 1
            continue

        # $$ 显示公式
        if part.startswith('$$') and part.endswith('$$'):
            inner = part[2:-2]
            try:
                latex_clean = clean_latex(inner)
                oMath = latex_to_omml(latex_clean)
                paragraph._element.append(oMath)
            except Exception:
                paragraph.add_run(f"[公式: {inner[:40]}]")

        # $ 行内公式
        elif part.startswith('$') and part.endswith('$'):
            inner = part[1:-1]
            try:
                latex_clean = clean_latex(inner)
                oMath = latex_to_omml(latex_clean)
                paragraph._element.append(oMath)
            except Exception:
                paragraph.add_run(f"[{inner[:30]}]")

        # **粗体**
        elif part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True

        # *斜体*
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.italic = True

        # `代码`
        elif part.startswith('`') and part.endswith('`'):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        # 链接 [text](url)
        elif len(parts) > i + 2 and part.startswith('['):
            # parts 结构: [full_match, text, url, ...]
            link_text = part[1:part.index(']')]
            i += 2  # skip the text and url capture groups
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True

        else:
            # 普通文本
            paragraph.add_run(part)

        i += 1


def process_inline_math(paragraph, text: str):
    """处理段落中的行内 $...$ 公式（旧接口，保持兼容）。"""
    process_markdown_text(paragraph, text)


def add_paragraph_text(doc: Document, text: str):
    """添加一个段落，处理标题、Markdown 语法、行内公式等。"""
    # 检查标题
    lev, content = heading_level(text)
    if lev is not None:
        p = doc.add_heading(content, level=lev)
        return p

    # 统一交给 process_markdown_text 处理（含 ** * ` $ $$ 等）
    p = doc.add_paragraph()
    process_markdown_text(p, text)
    return p


def add_mermaid_block(doc: Document, code: str):
    """添加 Mermaid 图的说明文字 + 源码。"""
    p = doc.add_paragraph()
    run = p.add_run("📊 Mermaid 图")
    run.bold = True
    run.font.size = Pt(11)

    # 提取第一个 graph 声明作为简单描述
    first_part = code.strip()[:80].replace('\n', ' | ')
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"示意图: {first_part}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 源码作为小字代码块
    p3 = doc.add_paragraph()
    run3 = p3.add_run(code)
    run3.font.size = Pt(8)
    run3.font.name = 'Courier New'
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(6)


def add_code_block(doc: Document, code: str):
    """添加代码块。"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'


def convert_md_to_docx(md_content: str, output_path: str,
                       md_source_dir: str = None):
    """将 Markdown 内容转换为 docx。"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10.5)

    # 解析块
    blocks = parse_blocks(md_content)

    for block in blocks:
        t = block['type']

        if t == 'paragraph':
            add_paragraph_text(doc, block['text'])

        elif t == 'table':
            add_table(doc, block['rows'])

        elif t == 'mermaid':
            add_mermaid_block(doc, block['code'])

        elif t == 'code':
            add_code_block(doc, block['code'])

    # 处理图片嵌入
    if md_source_dir:
        embed_images(doc, md_source_dir)

    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")


def embed_images(doc: Document, source_dir: str):
    """从 markdown 图片引用嵌入图片到 docx。"""
    # 搜索 ![](path) 模式
    img_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.svg', '.bmp']
    for p in doc.paragraphs:
        text = p.text
        for m in re.finditer(r'!\[([^\]]*)\]\(([^)]+)\)', text):
            img_path = m.group(2)
            # 尝试多种路径
            candidates = [
                img_path,
                os.path.join(source_dir, img_path),
                os.path.join(source_dir, 'imgs', img_path),
                os.path.join(source_dir, 'images', img_path),
                os.path.join(source_dir, 'img', img_path),
            ]
            for c in candidates:
                if os.path.isfile(c):
                    try:
                        run = p.add_run()
                        run.add_picture(c, width=Inches(4.5))
                    except Exception:
                        pass
                    break


# ═══════════════════════════════════════════════════
# 文件收集
# ═══════════════════════════════════════════════════


def collect_md_files(directory: str) -> list[str]:
    if not os.path.isdir(directory):
        print(f"❌ 目录不存在: {directory}", file=sys.stderr)
        sys.exit(1)
    files = []
    for f in sorted(os.listdir(directory)):
        if not f.endswith(".md"):
            continue
        if f == "README.md" or "_bak" in f or ".bak" in f:
            continue
        full = os.path.join(directory, f)
        if os.path.isfile(full):
            files.append(full)
    if not files:
        print(f"❌ 目录中没有 .md 文件: {directory}", file=sys.stderr)
        sys.exit(1)
    return files


# ═══════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════


def cmd_single(args):
    md_path = args.md_path
    if not os.path.isfile(md_path):
        print(f"❌ 文件不存在: {md_path}", file=sys.stderr)
        sys.exit(1)
    output = args.output or os.path.splitext(md_path)[0] + ".docx"
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()
    md_dir = os.path.dirname(os.path.abspath(md_path))
    convert_md_to_docx(content, output, md_source_dir=md_dir)


def cmd_dir(args):
    directory = args.directory
    files = collect_md_files(directory)
    output = args.output or os.path.join(directory,
        os.path.basename(os.path.normpath(directory)) + ".docx")
    print(f"📄 合并 {len(files)} 个 .md 文件 → {output}")

    all_content = []
    for i, fp in enumerate(files):
        print(f"   + {os.path.basename(fp)}")
        with open(fp, "r", encoding="utf-8") as f:
            content = f.read()
        # 剥离 frontmatter（非首文件也剥离）
        content = re.sub(r'^---\n.*?\n---\n', '', content, flags=re.DOTALL)
        if i > 0:
            content = "\n\n---\n\n" + content  # 章节分隔线
        all_content.append(content)

    md_content = "\n\n".join(all_content)
    convert_md_to_docx(md_content, output, md_source_dir=directory)


def main():
    p = argparse.ArgumentParser(
        description="教材 Markdown → Word 转换工具（v3，OMML 公式）"
    )
    sp = p.add_subparsers(dest="cmd", required=True)
    s = sp.add_parser("single", help="转换单个 .md 文件为 .docx")
    s.add_argument("md_path")
    s.add_argument("-o", "--output")
    d = sp.add_parser("dir", help="将目录下所有 .md 合并转换为一个 .docx")
    d.add_argument("directory")
    d.add_argument("-o", "--output")
    a = p.parse_args()
    if a.cmd == "single":
        cmd_single(a)
    elif a.cmd == "dir":
        cmd_dir(a)


if __name__ == "__main__":
    main()

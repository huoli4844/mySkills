#!/usr/bin/env python3
"""
大纲解析工具 — 从 .docx / .md / 纯文本中提取章节分层树。

用法：
  python3 parse_outline.py 大纲.docx -o /tmp/outline.json
  python3 parse_outline.py 大纲.md -o /tmp/outline.json
  cat 大纲.txt | python3 parse_outline.py /dev/stdin -o /tmp/outline.json
  python3 parse_outline.py /dev/stdin -o /tmp/outline.json   # 粘贴后 Ctrl+D

输出 JSON 结构：
{
  "title": "书名",
  "chapters": [
    {"number": "1", "title": "第1章 绪论", "sections": [
      {"number": "1.1", "title": "1.1 研究背景", "subsections": []},
      ...
    ]}
  ]
}
"""
import argparse, json, re, sys
from pathlib import Path


CHAPTER_PATTERN = re.compile(r'^(?:#{1,2}\s*)?第\s*([一二三四五六七八九十\d]+)\s*章\s*(.*?)$', re.MULTILINE)
SECTION_PATTERN = re.compile(r'^#{2,3}\s*(\d+\.\d+)\s*(.*?)$', re.MULTILINE)
SUBSECTION_PATTERN = re.compile(r'^#{3,4}\s*(\d+\.\d+\.\d+)\s*(.*?)$', re.MULTILINE)

CHINESE_NUM = {'一':'1','二':'2','三':'3','四':'4','五':'5',
               '六':'6','七':'7','八':'8','九':'9','十':'10'}


def chn_to_arabic(cn: str) -> str:
    if cn.isdigit():
        return cn
    if cn in CHINESE_NUM:
        return CHINESE_NUM[cn]
    if '十' in cn:
        if cn == '十': return '10'
        if cn.startswith('十'): return f'1{CHINESE_NUM.get(cn[1],"0")}'
        if cn.endswith('十'): return f'{CHINESE_NUM.get(cn[0],"1")}0'
        return f'{CHINESE_NUM.get(cn[0],"1")}{CHINESE_NUM.get(cn[1],"0")}'
    return cn


SEC_pat = re.compile(r'^\s*(\d+\.\d+)\s+(.*)')
SUB_pat = re.compile(r'^\s*(\d+\.\d+\.\d+)\s+(.*)')
PART_pat = re.compile(r'第\s*(\d+|[一二三四五六七八九十]+)\s*部分\s*(.*)')
CASE_pat = re.compile(r'案例\s*(\d+)')
EXP_pat = re.compile(r'实验\s*(\d+)')


def parse_docx(docx_path: str) -> dict:
    """从 .docx 解析大纲（支持无 Heading 样式的纯文本编号大纲）"""
    from docx import Document
    doc = Document(str(Path(docx_path).expanduser().resolve()))
    
    # 第一步：提取所有段落并分类
    entries = []  # [(type, number, title, paragraph_index)]
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        bold = any(r.bold for r in p.runs if r.bold)
        
        m = re.match(r'第\s*(\d+|[一二三四五六七八九十]+)\s*章\s*(.*)', text)
        if m:
            ch_num = m.group(1).strip()
            ch_title = m.group(2).strip()
            ch_arabic = chn_to_arabic(ch_num)
            entries.append(('chapter', ch_arabic, f'第{ch_arabic}章 {ch_title}', i))
            continue
        
        m = PART_pat.match(text)
        if m:
            pt_num = m.group(1).strip()
            pt_title = m.group(2).strip()
            pt_arabic = chn_to_arabic(pt_num)
            entries.append(('part', pt_arabic, f'第{pt_arabic}部分 {pt_title}', i))
            continue
        
        m = SUB_pat.match(text)
        if m:
            entries.append(('subsection', m.group(1), m.group(2).strip(), i))
            continue
        
        m = SEC_pat.match(text)
        if m:
            entries.append(('section', m.group(1), m.group(2).strip(), i))
            continue
        
        m = CASE_pat.match(text)
        if m:
            entries.append(('case', m.group(1), text, i))
            continue
        
        m = EXP_pat.match(text)
        if m:
            entries.append(('experiment', m.group(1), text, i))
            continue
    
    # 第二步：构建章节树
    chapters = []
    current_chapter = None
    current_section = None
    
    for entry in entries:
        t, num, title, idx = entry
        
        if t == 'chapter':
            current_chapter = {'number': num, 'title': title, 'sections': []}
            chapters.append(current_chapter)
            current_section = None
        
        elif t == 'section' and current_chapter is not None:
            current_section = {'number': num, 'title': title, 'subsections': []}
            current_chapter['sections'].append(current_section)
        
        elif t == 'subsection' and current_section is not None:
            current_section['subsections'].append({'number': num, 'title': title})
    
    # 书名：取第一个加粗的短段落（不含章/部分）
    title = ""
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text: continue
        bold = any(r.bold for r in p.runs if r.bold)
        if bold and len(text) <= 30 and '章' not in text and '部分' not in text:
            title = text
            break
    
    return {"title": title, "chapters": chapters}


def parse_text(text: str) -> dict:
    lines = text.strip().split('\n')
    title = ""
    for line in lines:
        s = line.strip().lstrip('#').strip()
        if s and not re.search(r'第.*章', s) and len(s) > 2:
            title = s
            break

    chapters = []
    for ch_match in CHAPTER_PATTERN.finditer(text):
        ch_num = ch_match.group(1).strip()
        ch_title_raw = ch_match.group(2).strip()
        ch_pos = ch_match.start()

        ch_arabic = chn_to_arabic(ch_num)
        ch_title = f"第{ch_arabic}章 {ch_title_raw}" if ch_title_raw else f"第{ch_arabic}章"

        next_ch = CHAPTER_PATTERN.search(text, ch_match.end())
        ch_end = next_ch.start() if next_ch else len(text)
        ch_text = text[ch_pos:ch_end]

        sections = []
        for sec_match in SECTION_PATTERN.finditer(ch_text):
            sec_num = sec_match.group(1)
            sec_title = sec_match.group(2).strip()
            sec_title = f"{sec_num} {sec_title}" if sec_title else sec_num

            sec_pos = sec_match.start()
            next_sec = SECTION_PATTERN.search(ch_text, sec_match.end())
            sec_end = next_sec.start() if next_sec else len(ch_text)
            sec_text = ch_text[sec_pos:sec_end]

            subsections = []
            for sub_match in SUBSECTION_PATTERN.finditer(sec_text):
                sub_num = sub_match.group(1)
                sub_title = sub_match.group(2).strip()
                sub_title = f"{sub_num} {sub_title}" if sub_title else sub_num
                subsections.append({"number": sub_num, "title": sub_title})

            sections.append({
                "number": sec_num,
                "title": sec_title,
                "subsections": subsections,
            })

        chapters.append({
            "number": ch_arabic,
            "title": ch_title,
            "sections": sections,
        })

    return {"title": title, "chapters": chapters}


def print_tasks(outline: dict):
    total_leaves = sum(len(c.get("sections", [])) for c in outline["chapters"])
    print(f"📋 写作任务清单（共计 {total_leaves} 个节）")
    for c in outline["chapters"]:
        secs = c.get("sections", [])
        print(f"  ├── {c['title']} ({len(secs)}节)")
        for s in secs:
            if s.get("subsections"):
                print(f"  │   ├── {s['title']} ({len(s['subsections'])}子节)")
                for sub in s["subsections"]:
                    print(f"  │   │   └── {sub['title']}")
            else:
                print(f"  │   └── {s['title']}")


def main():
    parser = argparse.ArgumentParser(description='大纲解析：提取章节分层树')
    parser.add_argument('input', help='输入文件路径（.docx / .md / /dev/stdin）')
    parser.add_argument('-o', '--output', required=True, help='输出 JSON 文件路径')
    parser.add_argument('--print-tasks', action='store_true', help='同时打印任务清单')
    args = parser.parse_args()

    input_path = Path(args.input)

    if args.input.endswith('.docx'):
        outline = parse_docx(args.input)
    else:
        text = input_path.read_text('utf-8', errors='replace') if input_path.exists() else sys.stdin.read()
        outline = parse_text(text)

    ch_nums = [int(c["number"]) for c in outline["chapters"] if c["number"].isdigit()]
    if ch_nums:
        expected = list(range(ch_nums[0], ch_nums[-1] + 1))
        missing = set(expected) - set(ch_nums)
        if missing:
            print(f"⚠️ 章节编号可能不连续，缺失: {missing}", file=sys.stderr)

    with open(args.output, 'w', encoding='utf-8') as f:
        json.dump(outline, f, ensure_ascii=False, indent=2)

    print(f"✅ 已输出: {args.output}（{len(outline['chapters'])}章，{sum(len(c.get('sections',[])) for c in outline['chapters'])}节）")

    if args.print_tasks:
        print()
        print_tasks(outline)


if __name__ == '__main__':
    main()
